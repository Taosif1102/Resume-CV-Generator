# Start Imports
from docx import Document
from docx.shared import Pt
# End Imports

def generate_resume():
    # Start Setup Document
    document = Document()
    # End Setup Document

    # Start Main Header
    # Main Header
    document.add_heading('Resume_Builder', 0)
    # End Main Header

    # Start Personal Information
    # Personal Information
    document.add_heading('Personal Information', level=1)
    document.add_paragraph('Name: [Your Name]')
    document.add_paragraph('Address: [Your Address]')
    document.add_paragraph('Phone: [Your Phone]')
    document.add_paragraph('Email: [Your Email]')
    # End Personal Information

    # Start Career Objective
    # Career Objective
    document.add_heading('Career Objective', level=1)
    document.add_paragraph(
        'To secure a challenging position in a reputable organization to expand my '
        'learnings, knowledge, and skills.'
    )
    # End Career Objective

    # Start Academic Details
    # Academic Details
    document.add_heading('Academic Details', level=1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Degree'
    hdr_cells[1].text = 'Institute'
    hdr_cells[2].text = 'Board/University'
    hdr_cells[3].text = 'Year'
    
    # Add a sample row
    row_cells = table.add_row().cells
    row_cells[0].text = 'SSC'
    row_cells[1].text = 'Sample School'
    row_cells[2].text = 'Sample Board'
    row_cells[3].text = '2020'
    # End Academic Details

    # Start Computer Literacy
    # Computer Literacy
    document.add_heading('Computer Literacy', level=1)
    document.add_paragraph('MS Word, MS Excel, PowerPoint, Python, HTML, CSS')
    # End Computer Literacy

    # Start Key Qualities
    # Key Qualities
    document.add_heading('Key Qualities', level=1)
    document.add_paragraph('Quick Learner', style='List Bullet')
    document.add_paragraph('Team Player', style='List Bullet')
    document.add_paragraph('Hardworking', style='List Bullet')
    # End Key Qualities

    # Start Save Document
    # Save the document
    output_filename = 'Resume_Builder_Output.docx'
    document.save(output_filename)
    print(f"Resume generated successfully: {output_filename}")
    # End Save Document

# Start Execution Block
if __name__ == "__main__":
    generate_resume()
# End Execution Block
