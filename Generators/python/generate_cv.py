# Start Imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
# End Imports

def create_cv():
    # Start Setup Document
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    # End Setup Document

    # Start Header
    # Header (Name and Title)
    name = document.add_paragraph()
    name_run = name.add_run('JOHN DOE')
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(50, 59, 76) # Dark Blue/Grey
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = document.add_paragraph()
    title_run = title.add_run('Creative Designer')
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(102, 102, 102) # Grey
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph() # Spacer
    # End Header

    # Start Table Layout
    # Create a table for the 2-column layout
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(5.0)
    # End Table Layout

    # Start Left Column
    # Left Column (Contact, Skills, Languages)
    left_cell = table.cell(0, 0)
    
    # Contact
    left_cell.add_paragraph('CONTACT').runs[0].bold = True
    left_cell.add_paragraph('Phone: +1 234 567 890')
    left_cell.add_paragraph('Email: john@example.com')
    left_cell.add_paragraph('Location: New York, USA')
    left_cell.add_paragraph('LinkedIn: linkedin.com/in/johndoe')
    
    left_cell.add_paragraph() # Spacer

    # Skills
    left_cell.add_paragraph('SKILLS').runs[0].bold = True
    skills = ['Graphic Design', 'UI/UX Design', 'Web Development', 'Adobe Suite', 'Figma']
    for skill in skills:
        left_cell.add_paragraph(f'• {skill}')
        
    left_cell.add_paragraph() # Spacer

    # Languages
    left_cell.add_paragraph('LANGUAGES').runs[0].bold = True
    languages = ['English (Native)', 'Spanish (Fluent)', 'French (Basic)']
    for lang in languages:
        left_cell.add_paragraph(f'• {lang}')
    # End Left Column

    # Start Right Column
    # Right Column (Profile, Experience, Education)
    right_cell = table.cell(0, 1)

    # Profile
    right_cell.add_paragraph('PROFILE').runs[0].bold = True
    right_cell.add_paragraph('Creative and detail-oriented designer with over 5 years of experience in creating visually stunning and user-friendly digital experiences. Passionate about branding, typography, and clean code.')
    
    right_cell.add_paragraph() # Spacer

    # Experience
    right_cell.add_paragraph('EXPERIENCE').runs[0].bold = True
    
    exp1 = right_cell.add_paragraph()
    exp1.add_run('Senior Designer').bold = True
    exp1.add_run('\nTech Solutions Inc. | 2020 - Present').italic = True
    right_cell.add_paragraph('• Led the redesign of the company website, increasing user engagement by 40%.\n• Mentored junior designers and established a new design system.\n• Collaborated with developers to ensure pixel-perfect implementation.')

    exp2 = right_cell.add_paragraph()
    exp2.add_run('Graphic Designer').bold = True
    exp2.add_run('\nCreative Agency | 2018 - 2020').italic = True
    right_cell.add_paragraph('• Designed marketing materials for various clients including logos, brochures, and social media graphics.\n• Managed multiple projects simultaneously while meeting tight deadlines.')

    right_cell.add_paragraph() # Spacer

    # Education
    right_cell.add_paragraph('EDUCATION').runs[0].bold = True
    
    edu1 = right_cell.add_paragraph()
    edu1.add_run('Bachelor of Fine Arts').bold = True
    edu1.add_run('\nUniversity of Design | 2014 - 2018').italic = True
    right_cell.add_paragraph('Specialized in Digital Arts and Multimedia.')
    # End Right Column

    # Start Save Document
    # Save the document
    output_path = 'docx/cv_builder_output.docx'
    document.save(output_path)
    print(f"CV generated successfully: {output_path}")
    # End Save Document

# Start Execution Block
if __name__ == "__main__":
    create_cv()
# End Execution Block
